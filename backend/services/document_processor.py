import json
import re
import uuid
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from backend.models.schemas import ProcessedDocument


def process_document(file_path: str, file_type: str) -> ProcessedDocument:
    path = Path(file_path)
    filename = path.name
    doc_id = str(uuid.uuid4())

    if file_type == "pdf":
        return _process_pdf(path, doc_id, filename)
    elif file_type == "docx":
        return _process_docx(path, doc_id, filename)
    elif file_type in ("txt", "md"):
        return _process_text(path, doc_id, filename, file_type)
    elif file_type == "html":
        return _process_html(path, doc_id, filename)
    elif file_type == "csv":
        return _process_csv(path, doc_id, filename)
    elif file_type == "json":
        return _process_json(path, doc_id, filename)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def process_text_content(text: str, title: str) -> ProcessedDocument:
    doc_id = str(uuid.uuid4())
    cleaned = _clean_text(text)
    return ProcessedDocument(
        id=doc_id,
        filename=title,
        file_type="txt",
        total_characters=len(cleaned),
        text_content=cleaned,
        sections=[{"title": title, "content": cleaned, "page": None}],
        metadata={"source": "pasted_text"},
    )


def process_html_content(html: str, url: str) -> ProcessedDocument:
    doc_id = str(uuid.uuid4())
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    title = soup.title.string if soup.title and soup.title.string else url

    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = _clean_text(main.get_text(separator="\n"))

    sections = []
    for heading in main.find_all(["h1", "h2", "h3"]):
        content = []
        for sib in heading.find_next_siblings():
            if sib.name in ("h1", "h2", "h3"):
                break
            content.append(sib.get_text(separator=" ").strip())
        sections.append(
            {
                "title": heading.get_text().strip(),
                "content": "\n".join(content),
                "page": None,
            }
        )

    if not sections:
        sections = [{"title": title, "content": text, "page": None}]

    return ProcessedDocument(
        id=doc_id,
        filename=title,
        file_type="html",
        total_characters=len(text),
        text_content=text,
        sections=sections,
        metadata={"url": url, "title": title},
    )


def _process_pdf(path: Path, doc_id: str, filename: str) -> ProcessedDocument:
    doc = fitz.open(str(path))
    pages_text = []
    sections = []
    for i, page in enumerate(doc):
        text = page.get_text()
        pages_text.append(text)
        sections.append({"title": f"Page {i + 1}", "content": text, "page": i + 1})
    doc.close()

    full_text = _clean_text("\n".join(pages_text))
    return ProcessedDocument(
        id=doc_id,
        filename=filename,
        file_type="pdf",
        total_characters=len(full_text),
        total_pages=len(pages_text),
        text_content=full_text,
        sections=sections,
        metadata={"page_count": len(pages_text)},
    )


def _process_docx(path: Path, doc_id: str, filename: str) -> ProcessedDocument:
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = _clean_text("\n\n".join(paragraphs))

    sections = []
    current_section = {"title": "Document", "content": [], "page": None}
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            if current_section["content"]:
                current_section["content"] = "\n".join(current_section["content"])
                sections.append(current_section)
            current_section = {"title": para.text, "content": [], "page": None}
        elif para.text.strip():
            current_section["content"].append(para.text)

    if current_section["content"]:
        current_section["content"] = "\n".join(current_section["content"])
        sections.append(current_section)

    if not sections:
        sections = [{"title": filename, "content": full_text, "page": None}]

    return ProcessedDocument(
        id=doc_id,
        filename=filename,
        file_type="docx",
        total_characters=len(full_text),
        text_content=full_text,
        sections=sections,
        metadata={},
    )


def _process_text(
    path: Path, doc_id: str, filename: str, file_type: str
) -> ProcessedDocument:
    text = path.read_text(encoding="utf-8")
    cleaned = _clean_text(text)
    return ProcessedDocument(
        id=doc_id,
        filename=filename,
        file_type=file_type,
        total_characters=len(cleaned),
        text_content=cleaned,
        sections=[{"title": filename, "content": cleaned, "page": None}],
        metadata={},
    )


def _process_html(path: Path, doc_id: str, filename: str) -> ProcessedDocument:
    html = path.read_text(encoding="utf-8")
    return process_html_content(html, filename)


def _process_csv(path: Path, doc_id: str, filename: str) -> ProcessedDocument:
    df = pd.read_csv(str(path))

    csv_columns = []
    for col in df.columns:
        dtype = str(df[col].dtype)
        sample_values = df[col].dropna().head(5).tolist()
        col_type = "text"
        if "int" in dtype or "float" in dtype:
            col_type = "numeric"
        elif "datetime" in dtype:
            col_type = "datetime"
        csv_columns.append(
            {"name": col, "type": col_type, "sample_values": sample_values}
        )

    text_repr = df.to_string(index=False)
    return ProcessedDocument(
        id=doc_id,
        filename=filename,
        file_type="csv",
        total_characters=len(text_repr),
        text_content=text_repr,
        sections=[],
        metadata={"row_count": len(df), "column_count": len(df.columns)},
        csv_columns=csv_columns,
    )


def _process_json(path: Path, doc_id: str, filename: str) -> ProcessedDocument:
    data = json.loads(path.read_text(encoding="utf-8"))
    text = json.dumps(data, indent=2)

    flat = _flatten_json(data) if isinstance(data, dict) else str(data)
    return ProcessedDocument(
        id=doc_id,
        filename=filename,
        file_type="json",
        total_characters=len(text),
        text_content=text if isinstance(text, str) else str(flat),
        sections=[{"title": filename, "content": text, "page": None}],
        metadata={"type": type(data).__name__},
    )


def _flatten_json(data: dict, prefix: str = "") -> dict:
    items = {}
    for k, v in data.items():
        key = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict):
            items.update(_flatten_json(v, key))
        elif isinstance(v, list):
            items[key] = str(v)
        else:
            items[key] = v
    return items


def _clean_text(text: str) -> str:
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines).strip()
